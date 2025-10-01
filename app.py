import io
import re
import unicodedata
from typing import Dict, List, Any, Iterable

import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

# =====================================================
# Utilidades
# =====================================================

PH_RE = re.compile(r"\{\{([^}]+)\}\}")

def strip_accents(s: str) -> str:
    import unicodedata
    return "".join(c for c in unicodedata.normalize("NFD", s) if unicodedata.category(c) != "Mn")

def normalize_key(s: str) -> str:
    s = strip_accents(s)
    s = re.sub(r"[^A-Za-z0-9]+", "_", s)
    return s.strip("_").upper()

def iter_all_paragraphs(doc: Document) -> Iterable:
    # corpo
    for p in doc.paragraphs:
        yield p
    # tabelas
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p

def collect_placeholders(doc: Document) -> List[str]:
    found = set()
    # corpo + tabelas
    for p in iter_all_paragraphs(doc):
        for ph in PH_RE.findall(p.text or ""):
            found.add(ph.strip())
    # headers/footers
    for sec in doc.sections:
        for hf in [sec.header, sec.first_page_header, sec.even_page_header,
                   sec.footer, sec.first_page_footer, sec.even_page_footer]:
            if hf:
                for p in hf.paragraphs:
                    for ph in PH_RE.findall(p.text or ""):
                        found.add(ph.strip())
    return sorted(found, key=lambda x: x.lower())

def replace_in_paragraph(p, mapping: Dict[str, Any]) -> bool:
    """Substitui placeholders respeitando runs quebrados."""
    full = "".join(run.text for run in p.runs)
    if not full:
        return False
    placeholders = PH_RE.findall(full)
    if not placeholders:
        return False
    new_full = full
    for raw in placeholders:
        raw_key = raw.strip()
        # tenta exatamente e a versão normalizada
        val = mapping.get(raw_key)
        if val is None:
            val = mapping.get(normalize_key(raw_key))
        if val is None:
            val = ""
        new_full = new_full.replace("{{" + raw_key + "}}", str(val))
    if new_full != full:
        if not p.runs:
            p.add_run(new_full)
        else:
            p.runs[0].text = new_full
            for r in p.runs[1:]:
                r.text = ""
        return True
    return False

def paragraph_contains_any(p, needles: List[str]) -> bool:
    t = p.text or ""
    return any((n.strip().lower() in t.lower()) for n in needles if n.strip())

def apply_font_to_paragraph(p, font_name: str, size_pt: int, bold: bool):
    for run in p.runs:
        run.font.name = font_name
        rPr = run._element.rPr
        if rPr is not None and rPr.rFonts is not None:
            rPr.rFonts.set(qn('w:ascii'), font_name)
            rPr.rFonts.set(qn('w:hAnsi'), font_name)
            rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = Pt(size_pt)
        run.font.bold = bold

def process_document(template_bytes: bytes,
                     mapping: Dict[str, Any],
                     font_name: str,
                     font_size: int,
                     bold_all: bool,
                     exceptions_contains: List[str],
                     exceptions_placeholders: List[str]) -> bytes:
    """Retorna bytes do DOCX final."""
    doc = Document(io.BytesIO(template_bytes))

    # Replace em corpo + tabelas
    for p in iter_all_paragraphs(doc):
        replace_in_paragraph(p, mapping)

    # Replace em header/footer
    for sec in doc.sections:
        for hf in [sec.header, sec.first_page_header, sec.even_page_header,
                   sec.footer, sec.first_page_footer, sec.even_page_footer]:
            if hf:
                for p in hf.paragraphs:
                    replace_in_paragraph(p, mapping)

    # Tipografia + regra de negrito
    exc_ph_norm = {normalize_key(x) for x in exceptions_placeholders}
    for p in iter_all_paragraphs(doc):
        is_exc_text = paragraph_contains_any(p, exceptions_contains)
        tnorm = normalize_key(p.text or "")
        is_exc_ph = any(k in tnorm for k in exc_ph_norm)
        apply_font_to_paragraph(p, font_name, font_size, bold_all and not (is_exc_text or is_exc_ph))

    for sec in doc.sections:
        for hf in [sec.header, sec.first_page_header, sec.even_page_header,
                   sec.footer, sec.first_page_footer, sec.even_page_footer]:
            if hf:
                for p in hf.paragraphs:
                    is_exc_text = paragraph_contains_any(p, exceptions_contains)
                    tnorm = normalize_key(p.text or "")
                    is_exc_ph = any(k in tnorm for k in exc_ph_norm)
                    apply_font_to_paragraph(p, font_name, font_size, bold_all and not (is_exc_text or is_exc_ph))

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.getvalue()

def safe_filename(name: str) -> str:
    return re.sub(r'[\\/:*?"<>|]+', "-", str(name)).strip()

# =====================================================
# UI
# =====================================================

st.set_page_config(page_title="Fornecedores Homologados — Automação", page_icon="🧩", layout="wide")
st.title("Carta de Homologação de Fornecedores — Preenchimento por {{Chaves}}")

st.markdown("**Fluxo:** ① Suba o modelo `.docx` → ② O app detecta as `{{CHAVES}}` → ③ Você preenche → ④ Baixe o `.docx` (e opcional `.pdf`).")
st.divider()

# 1) Upload do modelo
template_file = st.file_uploader("Modelo .docx (com placeholders {{CHAVE}})", type=["docx"], accept_multiple_files=False)

# 2) Config de tipografia
with st.sidebar:
    st.header("Configurações")
    font_name = st.text_input("Fonte", value="Calibri")
    font_size = st.number_input("Tamanho (pt)", 8, 20, 11, 1)
    bold_all = st.checkbox("Negrito por padrão (todo documento)", value=True)
    try_pdf = st.checkbox("Tentar exportar PDF (pode não funcionar no Cloud)", value=False)
    st.caption("PDF depende de Word/LibreOffice no host; em alguns ambientes não estará disponível.")

    st.markdown("---")
    st.subheader("Exceções (não negritar)")
    exc_contains = st.text_area("Exceções por **texto contido** no parágrafo (1 por linha):",
                                value="Descrever aqui os produtos ou serviços aprovados.").splitlines()
    exc_placeholders = st.text_area("Exceções por **placeholder** (a chave cujos valores não ficam em negrito, 1 por linha):",
                                    value="DESCRICAO_PRODUTOS_SERVICOS").splitlines()

# 3) Detectar chaves e montar formulário
if template_file:
    try:
        tmp = Document(template_file)
        placeholders = collect_placeholders(tmp)
        st.success(f"Placeholders detectados ({len(placeholders)}): {placeholders if placeholders else '—'}")

        with st.form("form_preenchimento"):
            st.subheader("Preencha os valores das {{Chaves}}")
            st.caption("Dica: campos em branco serão substituídos por vazio.")
            values: Dict[str, Any] = {}

            col1, col2 = st.columns(2)
            half = (len(placeholders) + 1) // 2
            left_keys = placeholders[:half]
            right_keys = placeholders[half:]

            with col1:
                for k in left_keys:
                    values[k] = st.text_input(k, value="")
            with col2:
                for k in right_keys:
                    values[k] = st.text_input(k, value="")

            output_name = st.text_input("Nome do arquivo de saída (.docx)",
                                        value="Homologacao - {RAZAO_SOCIAL_DO_FORNECEDOR}.docx")
            submitted = st.form_submit_button("Gerar documento")

        if submitted:
            # Monta mapping com chaves originais e normalizadas
            mapping: Dict[str, Any] = {}
            for k, v in values.items():
                mapping[k] = v
                mapping[normalize_key(k)] = v

            # Nome do arquivo (substitui tokens por valores preenchidos também)
            final_name = output_name
            for k, v in values.items():
                final_name = final_name.replace("{"+k+"}", safe_filename(v))
                final_name = final_name.replace("{"+normalize_key(k)+"}", safe_filename(v))
            if not final_name.lower().endswith(".docx"):
                final_name += ".docx"
            final_name = safe_filename(final_name) or "Homologacao.docx"

            # Recarrega bytes do modelo (file_uploader avança o ponteiro)
            template_file.seek(0)
            template_bytes = template_file.read()

            # Processa DOCX
            out_docx_bytes = process_document(
                template_bytes=template_bytes,
                mapping=mapping,
                font_name=font_name,
                font_size=int(font_size),
                bold_all=bold_all,
                exceptions_contains=[s for s in exc_contains if s.strip()],
                exceptions_placeholders=[s for s in exc_placeholders if s.strip()],
            )

            st.download_button("⬇️ Baixar DOCX", data=out_docx_bytes, file_name=final_name, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

            # PDF opcional
            if try_pdf:
                try:
                    # Tentativa de docx2pdf local (pode falhar no Cloud)
                    from docx2pdf import convert  # noqa: F401
                    import tempfile, os
                    with tempfile.TemporaryDirectory() as td:
                        tmp_docx = f"{td}/tmp.docx"
                        tmp_pdf = f"{td}/tmp.pdf"
                        with open(tmp_docx, "wb") as f:
                            f.write(out_docx_bytes)
                        # convert precisa de caminho; saída direta
                        convert(tmp_docx, tmp_pdf)  # type: ignore
                        with open(tmp_pdf, "rb") as f:
                            pdf_bytes = f.read()
                    pdf_name = final_name.replace(".docx", ".pdf")
                    st.download_button("⬇️ Baixar PDF", data=pdf_bytes, file_name=pdf_name, mime="application/pdf")
                except Exception as e:
                    st.info(f"Não foi possível gerar PDF neste ambiente: {e}")

    except Exception as e:
        st.error(f"Falha ao ler o modelo: {e}")
else:
    st.info("Envie o modelo .docx para começar.")
